import json, os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from enum import Enum
from PIL import Image
import pandas as pd

class PageSize(Enum):
    A3 = (11.69, 16.54)
    A4 = (8.27, 11.69)
    A5 = (5.83, 8.27)
    LETTER = (8.5, 11)
    LEGAL = (8.5, 14)
    TABLOID = (11 , 17)
    
    @property
    def dimensions(self):
        return self.value

class Orientation(Enum):
    PORTRAIT = 'portrait'
    LANDSCAPE = 'landscape'
    
    @property
    def Orient(self):
        return self.value
   
class TableStyle:
    def __init__(self, style:str = 'Colorful List Accent 6'):
        self.style = style
        
        
class FontStyle:
    def __init__(self, style:str='Normal', font_name:str = 'Segoe UI', font_size:int= 13):
        self.style = style
        self.fontName = font_name
        self.fontSize = Pt(font_size)


class ExportToWord:
    
    def __init__(self, output_file:str='word_output.docx', 
                 font_style: FontStyle=FontStyle(),
                 table_style: TableStyle=TableStyle(),
                 paper_size:PageSize= PageSize.A4, 
                 orientation:Orientation = Orientation.PORTRAIT):        
        self.paperSize = paper_size        
        self.Orientation = orientation
        self.fontStyle = font_style
        self.tableStyle = table_style
        self.outputFile = output_file
        self.pageDimension = paper_size.dimensions
        
    def setPageSize(self, section):
        width, height = self.paperSize.dimensions
        
        if self.Orientation == Orientation.LANDSCAPE:
            section.page_width = Inches(height)
            section.page_height = Inches(width)
            self.pageDimension = (height, width)
        else:
            section.page_width = Inches(width)
            section.page_height = Inches(height)
            self.pageDimension = (width, height)
    
    # def setFontStyle(self, style):        
    #     font = style.font
    #     font.name = self.fontStyle.fontName
    #     font.size = self.fontStyle.fontSize    
    #     return font
    
    def generate(self, doc_content:str):
        """Generate a Word document from the provided content with formatting like HTML.
        
        Params:
            doc_content (str): Document content to format.
        Returns:
            output_file (str): Path to save the generated Word document.
        """
        soup = BeautifulSoup(doc_content, 'html.parser')
        doc = Document()
        section = doc.sections[0]
        self.setPageSize(section)
        
        # Set default font and size for document
        style = doc.styles[self.fontStyle.style]
        font = style.font
        font.name = self.fontStyle.fontName
        font.size = self.fontStyle.fontSize
        # self.setFontStyle(style)
        
        # Process header and footer
        header = soup.find('header')
        footer = soup.find('footer')
        if header:
            self._formatHeaderFooter(section.header, header, cols=3)
        if footer:
            self._formatHeaderFooter(section.footer, footer)
    
        for element in soup.contents:
            if element.name is None:
                doc.add_paragraph(element.string.strip())
            elif element.name in ['header', 'footer']:
                pass
            elif element.name == 'title':
                self.addTitle(doc, element.text.strip())
            elif element.name in ['h1','h2','h3','h4','h5','h6']:
                self.addHeading(doc, element.text.strip(), int(element.name[1]))
            elif element.name == 'dataframe':
                # Process dataframe element
                if 'src' in element.attrs:
                    csv_file = element.get('src')
                    if csv_file:
                        df = pd.read_csv(csv_file)
                        self.addTable(doc, df)
                elif element.string and element.string.strip():
                    # Table is dict data as JSON string
                    _text = element.string.strip()
                    try:
                        _json = json.loads(_text)
                        df = pd.DataFrame.from_dict(_json)
                        self.addTable(doc, df)
                    except json.JSONDecodeError:
                        doc.add_paragraph(f"[Invalid JSON data]:{_text}")
            elif element.name == 'image':
                # Process image element
                if 'src' in element.attrs:
                    _img_file = element.get('src')
                    if _img_file:
                        self.addImage(doc, _img_file)
                    elif element.string and element.string.strip():
                        _base64 = element.string.strip()
                        _img_path = self.convertBase64ToImage(_base64)
                        if _img_path:
                            self.addImage(doc, _img_path)
            elif element.string and element.string.strip():
                doc.add_paragraph(element.string.strip())

        #Save document to output path
        doc.save(self.outputFile)
    
    def _formatHeaderFooter(self, section_part, element, rows=1, cols=1, height=0.5):
        """Process the header or footer of the Word document
        
        Args:
            section (docx.SectionPart): Header or footer of the section            
            element (Tag): BeautifulSoup tag for header or footer content.
            rows (int): rows of header or footer
            cols (int): columns of header or footer
        """        
        _rows = rows
        _cols = cols
        # _table_width = int(self.pageDimension[0]/_cols)
        _table_width = Inches(self.pageDimension[0] / _cols)
        table = section_part.add_table(rows=_rows, cols=_cols, width=Inches(self.pageDimension[0]))
        table.autofit = True                        
        table.alignment = WD_ALIGN_VERTICAL.CENTER
        
        for column in table.columns:
            column.width = _table_width
        
        first_cell = table.rows[0].cells[0]
        first_cell.text = element.text.strip()
                
        last_cell = table.rows[0].cells[-1]
        _image = element.find('image')
        if _image and 'src' in _image.attrs:
            _image_path = _image['src']
            paragraph = last_cell.paragraphs[0]
            if os.path.exists(_image_path):
                _width, _height = self.getImageSize(_image_path, height)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run()
                run.add_picture(_image_path, width = Inches(_width), height = Inches(_height))
            else:                
                run = paragraph.add_run(f'[{_image_path}]')
                run.font.color.rgb = RGBColor(255,0,0)
                run.bold = True

    def getImageSize(self, image_path, target_height=0.0):
        """Calculates the target size for the image while preserving the aspect ratio.

        Args:
            image_path (str): Path to the image file.
            target_height (float): Desired height of the image in inches.

        Returns:
            tuple: Width and height of the image in inches.
        """
        with Image.open(image_path) as img:
            img_width, img_height = img.size
            aspect_ratio  = img_width/img_height
            # target_height = img_height*0.00833 if target_height==0.0 else target_height
            if target_height == 0.0:
                target_width = self.pageDimension[0]*0.6
                target_height = target_width/aspect_ratio
            else:
                target_width = target_height*aspect_ratio
        return (target_width, target_height)
    
    def addTitle(self, doc, text):
        """Add a title to the Word document with custom formatting.
        
        Args:
            doc (Document): Word document object.
            text (str): Text for the title.
        """
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text.upper())  # Title should be uppercase
        run.bold = True
        # run.font.size = self.fontStyle.fontSize
        # run.font.name = self.fontStyle.fontName

    def addHeading(self, doc, text, level):
        """Add a heading to the Word document with custom formatting.
        
        Args:
            doc (Document): Word document object.
            text (str): Text for the heading.
            level (int): Heading level (1-6).
        """
        heading = doc.add_heading(level=level).add_run(text)  # Keep the text case as is
        # heading.font.name = self.fontStyle.fontName
        # heading.font.size = self.fontStyle.fontSize
        paragraph = heading._element.getparent().getparent()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def addTable(self, doc, dataframe):
        """Add a dataframe as a list table to a Word document with specific formatting.
        
        Args:
            doc (Document): Word document object.
            dataframe (pd.DataFrame): DataFrame to add as a table.
        """
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = self.tableStyle.style  # Apply the specific style
        
        # Add the header row
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(dataframe.columns):
            hdr_cells[i].text = column
            hdr_cells[i].paragraphs[0].runs[0].bold = True  # Make header bold
        
        # Add the data rows
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].runs[0].bold = False #make first cell of row normal
                

    def addImage(self, doc, image_path):
        """Add an image to the Word document, or replace with a placeholder if the file is missing.
        
        Args:
            doc (Document): Word document object.
            image_path (str): Path to the image file.
        """        
        if os.path.exists(image_path):            
            _width, _height = self.getImageSize(image_path)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(_width), height=Inches(_height))  # Adjust the image width
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"[Image not found: {image_path}]")
            run.font.color.rgb = RGBColor(255,0,0)
            run.bold = True
            
    def convertBase64ToImage(self, base64_string, output_file:str = "temp_image.png"):
        """Convert a Base64 string to an image file.
        
        Args:
            base64 (str): Base64 string representing image data with or without the data URL prefix.
            output_file (str): Path to save the converted image.
        
        Returns:
            str: Path to the saved image file.
        """
        # Remove the data URL prefix if present
        if base64_string.startswith("data:image"):
            base64 = base64_string.split(",")[1]  # Split at the comma and take the Base64 part

        # Decode the Base64 data and save it as an image file        
        with open(output_file, "wb") as f:
            f.write(base64.b64decode(base64))
        return output_file
    